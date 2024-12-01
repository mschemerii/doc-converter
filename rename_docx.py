#!/usr/bin/env python3
import sys
import os
import shutil
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def add_centered_header(doc, header_text):
    """Add centered header to the document"""
    # Get or create the header
    section = doc.sections[0]
    header = section.header
    
    # Clear any existing content in the header
    for paragraph in header.paragraphs:
        p = paragraph._p
        p.getparent().remove(p)
    
    # Create new paragraph in header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    
    # Add text and center it
    header_paragraph.text = header_text
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some spacing after the header
    pPr = header_paragraph._p.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="0" w:line="240" w:lineRule="auto"/>')
    pPr.append(spacing)
    
    return doc

def find_section_paragraph(doc, section_names):
    """Find paragraph that matches any of the given section names (case-insensitive)"""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip().lower()
        if any(name.lower() in text for name in section_names):
            return paragraph._p
    return None

def remove_between_sections(doc, start_section_names, end_section_names):
    """Remove content between start and end sections (inclusive for start, exclusive for end)"""
    # Find the paragraphs containing start and end sections
    start_para = find_section_paragraph(doc, start_section_names)
    end_para = find_section_paragraph(doc, end_section_names)
    
    if not start_para or not end_para:
        print(f"Warning: Could not find one or both sections: {start_section_names} to {end_section_names}")
        return doc
    
    # Get the parent element and indices
    parent = start_para.getparent()
    if not parent:
        return doc
    
    start_index = parent.index(start_para)
    end_index = parent.index(end_para)
    
    if start_index >= end_index:
        print(f"Warning: Start section appears after end section")
        return doc
    
    # Get elements between sections (inclusive start, exclusive end)
    elements_to_remove = []
    for child in parent:
        index = parent.index(child)
        if start_index <= index < end_index:  # Note: inclusive start, exclusive end
            elements_to_remove.append(child)
    
    # Remove the elements
    for element in elements_to_remove:
        parent.remove(element)
    
    return doc

def remove_from_section_to_end(doc, section_names):
    """Remove content from a section to the end of document (inclusive)"""
    section_para = find_section_paragraph(doc, section_names)
    
    if not section_para:
        print(f"Warning: Could not find section: {section_names}")
        return doc
    
    # Get the parent element and index
    parent = section_para.getparent()
    if not parent:
        return doc
    
    section_index = parent.index(section_para)
    
    # Remove everything from section to end (inclusive)
    elements_to_remove = []
    for child in parent:
        if parent.index(child) >= section_index:
            elements_to_remove.append(child)
    
    # Remove the elements
    for element in elements_to_remove:
        parent.remove(element)
    
    return doc

def create_renamed_copy(docx_path, suffix, header_text, remove_sections=None):
    """
    Create a copy of the .docx file with specific suffix and header:
    1. Replace +-+ with _
    2. Remove + characters to join text directly
    3. Add specified suffix before .docx extension
    4. Add centered header text
    5. Apply specified section removals
    """
    # Get the directory and filename
    directory = os.path.dirname(docx_path)
    filename = os.path.basename(docx_path)
    
    # Process the filename
    # First replace +-+ with _
    new_filename = filename.replace("+-+", "_")
    # Then remove remaining + characters
    new_filename = new_filename.replace("+", "")
    
    # Add suffix before extension
    base_name, ext = os.path.splitext(new_filename)
    new_filename = f"{base_name}-{suffix}{ext}"
    
    # Create the new file path
    new_path = os.path.join(directory, new_filename)
    
    try:
        # First copy the file
        shutil.copy2(docx_path, new_path)
        
        # Open the copy and add header
        doc = Document(new_path)
        doc = add_centered_header(doc, header_text)
        
        # Apply section removals if specified
        if remove_sections:
            for action, start_section, end_section in remove_sections:
                if action == 'remove_between':
                    doc = remove_between_sections(doc, start_section, end_section)
                elif action == 'remove_from_to_end':
                    doc = remove_from_section_to_end(doc, start_section)
        
        doc.save(new_path)
        
        print(f"Successfully created renamed copy with header: {new_filename}")
        return new_path
    except Exception as e:
        print(f"Error creating copy: {str(e)}")
        return None

def create_all_copies(docx_path):
    """Create three copies with different suffixes and headers"""
    versions = [
        # Stage-Evidence: Remove from Rollback to end
        (
            "Stage-Evidence",
            "Deploy to Stage",
            [('remove_from_to_end', ['Rollback'], None)]
        ),
        # StageDR-Evidence: Remove from Rollback to end
        (
            "StageDR-Evidence",
            "Deploy to StageDR",
            [('remove_from_to_end', ['Rollback'], None)]
        ),
        # Rollback-Evidence: Remove Pre-Deploy Steps up to (but not including) Rollback
        (
            "Rollback-Evidence",
            "Rollback",
            [('remove_between', ['Pre-Deploy Steps', 'Pre Steps'], ['Rollback'])]
        )
    ]
    
    created_files = []
    for suffix, header, remove_sections in versions:
        new_path = create_renamed_copy(docx_path, suffix, header, remove_sections)
        if new_path:
            created_files.append(new_path)
    
    return created_files

def main():
    if len(sys.argv) != 2:
        print("Usage: python rename_docx.py <path_to_docx_file>")
        sys.exit(1)
    
    try:
        docx_path = sys.argv[1]
        if not os.path.exists(docx_path):
            print(f"Error: File not found: {docx_path}")
            sys.exit(1)
        
        print(f"Original file: {docx_path}")
        created_files = create_all_copies(docx_path)
        
        if created_files:
            print("\nCreated the following files:")
            for file_path in created_files:
                print(f"- {os.path.basename(file_path)}")
    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)

if __name__ == '__main__':
    main()
