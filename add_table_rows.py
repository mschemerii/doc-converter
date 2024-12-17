#!/usr/bin/env python3
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from copy import deepcopy
import xml.etree.ElementTree as ET

def has_content(row):
    """Check if a row has any non-empty content"""
    return any(cell.text.strip() for cell in row.cells)

def copy_xml_element(element):
    """Create a copy of an XML element"""
    if element is None:
        return None
    return parse_xml(element.xml)

def copy_row_formatting(source_row, target_row):
    """Copy formatting from source row to target row"""
    # Copy row height if it exists
    source_tr = source_row._tr
    target_tr = target_row._tr
    
    if source_tr.trPr is not None:
        # Create new trPr element
        new_trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>') 
        
        # Copy height if it exists
        height = source_tr.trPr.find('.//w:trHeight', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if height is not None:
            new_height = copy_xml_element(height)
            new_trPr.append(new_height)
        
        # Replace or add to target
        if target_tr.trPr is not None:
            target_tr.remove(target_tr.trPr)
        target_tr.insert(0, new_trPr)
    
    # Copy cell formatting for each cell
    for source_cell, target_cell in zip(source_row.cells, target_row.cells):
        source_tc = source_cell._tc
        target_tc = target_cell._tc
        
        if source_tc.tcPr is not None:
            # Create new tcPr element
            new_tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>') 
            
            # Copy width if it exists
            width = source_tc.tcPr.find('.//w:tcW', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if width is not None:
                new_width = copy_xml_element(width)
                new_tcPr.append(new_width)
            
            # Copy vertical merge if it exists
            vmerge = source_tc.tcPr.find('.//w:vMerge', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if vmerge is not None:
                new_vmerge = copy_xml_element(vmerge)
                new_tcPr.append(new_vmerge)
            
            # Replace or add to target
            if target_tc.tcPr is not None:
                target_tc.remove(target_tc.tcPr)
            target_tc.insert(0, new_tcPr)
        
        # Copy paragraph formatting
        if source_cell.paragraphs and target_cell.paragraphs:
            source_p = source_cell.paragraphs[0]._p
            target_p = target_cell.paragraphs[0]._p
            
            if source_p.pPr is not None:
                new_pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>') 
                
                # Copy alignment if it exists
                jc = source_p.pPr.find('.//w:jc', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                if jc is not None:
                    new_jc = copy_xml_element(jc)
                    new_pPr.append(new_jc)
                
                # Replace or add to target
                if target_p.pPr is not None:
                    target_p.remove(target_p.pPr)
                target_p.insert(0, new_pPr)

def add_merged_row_after(table, row_index):
    """Add a new row after a specific row, with all cells merged."""
    current_row = table.rows[row_index]
    new_row = table.add_row()  # Add a new row

    # Clear text in the new row
    for cell in new_row.cells:
        cell.text = ""

    # Merge all cells in the new row
    if len(new_row.cells) > 1:
        new_row.cells[0].merge(new_row.cells[-1])

def add_rows_to_tables(docx_path):
    """
    Add a new empty row after each content row in all tables
    """
    # Load the document
    doc = Document(docx_path)
    
    # Process each table in the document
    for table in doc.tables:
        # Keep track of how many rows we've added to adjust indices
        added_rows = 0
        
        # Iterate through original rows
        for i in range(len(table.rows)):
            # Adjust index for added rows
            current_row = table.rows[i + added_rows]
            
            # If this row has content, add a new row after it
            if has_content(current_row):
                # Add new row
                new_row = table.add_row()
                
                # Move the new row to the position after the current row
                table._tbl.remove(new_row._tr)
                table._tbl.insert(current_row._tr.getparent().index(current_row._tr) + 1, new_row._tr)
                
                # Copy formatting from the content row to the new row
                copy_row_formatting(current_row, new_row)
                
                # Increment our counter
                added_rows += 1
    
    # Save the modifications back to the same file
    doc.save(docx_path)

def process_document(doc):
    """Process tables to delete empty rows/tables and add merged rows."""
    tables = doc.tables
    i = 0

    while i < len(tables):
        table = tables[i]
        first_cell_text = table.rows[0].cells[0].text.strip() if table.rows else ""

        # Skip specific tables
        if "Change request numbers" in first_cell_text or "Manifests" in first_cell_text or "Sync Multi Manifest String" in first_cell_text:
            i += 1
            continue

        # Process "Task No." tables without adding a row after the first row
        rows_added = 0
        for j in range(len(table.rows)):
            adjusted_index = j + rows_added
            if has_content(table.rows[adjusted_index]):
                add_merged_row_after(table, adjusted_index)
                rows_added += 1

        i += 1

def main():
    if len(sys.argv) != 2:
        print("Usage: python add_table_rows.py <path_to_docx_file>")
        sys.exit(1)
    
    try:
        docx_path = sys.argv[1]
        doc = Document(docx_path)
        process_document(doc)
        doc.save(docx_path)
        print(f"Successfully added rows to tables in: {docx_path}")
    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)

if __name__ == '__main__':
    main()
